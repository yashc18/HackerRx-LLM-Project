import asyncio
import hashlib
import logging
import time
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import io

import aiohttp
import fitz  # PyMuPDF
import docx
import email
from email import policy
import pytesseract
from pdf2image import convert_from_bytes
import numpy as np
from PIL import Image
import re

from utils.config import Settings

logger = logging.getLogger(__name__)
settings = Settings()

class AdvancedDocumentProcessor:
    def __init__(self, cache=None, settings=None):
        self.cache = cache
        self.settings = settings or Settings()
        self.quality_thresholds = {
            'min_text_length': self.settings.min_text_length,
            'max_processing_time': self.settings.max_processing_time_seconds,
            'ocr_confidence': 0.7,  # Default value
            'extraction_success_rate': 0.8  # Default value
        }
        self.supported_formats = {'.pdf', '.docx', '.eml', '.txt', '.html'}
        
    async def process_document(self, url: str) -> Dict[str, Any]:
        """Simple document processing method for compatibility"""
        return await self.process_with_robustness(url)
    
    async def process_with_robustness(self, url: str) -> Dict[str, Any]:
        """Robust document processing with comprehensive error handling"""
        start_time = time.time()
        
        try:
            # Download with retry mechanism
            content = await self.download_with_retries(url, max_retries=3)
            
            # Multi-format processing with fallbacks
            if url.lower().endswith('.pdf'):
                result = await self.process_pdf_advanced(content)
            elif url.lower().endswith('.docx'):
                result = await self.process_docx_advanced(content)
            elif url.lower().endswith('.eml'):
                result = await self.process_email_advanced(content)
            else:
                result = await self.process_generic_text(content)
            
            # Quality gates for hackathon evaluation
            quality_score = self.validate_extraction_quality(result)
            if quality_score < 0.8:  # Default quality threshold
                result = await self.enhance_extraction(result, content)
            
            # Add metadata for evaluation
            result['processing_metadata'] = {
                'processing_time': time.time() - start_time,
                'quality_score': quality_score,
                'extraction_confidence': self.calculate_extraction_confidence(result),
                'format_detected': self.detect_document_format(url, content),
                'total_sections': len(result.get('sections', [])),
                'image_count': len(result.get('images', [])),
                'table_count': len(result.get('tables', [])),
                'url': url
            }
            
            return result
            
        except Exception as e:
            return self.handle_processing_error(e, url)
    
    async def download_with_retries(self, url: str, max_retries: int = 3) -> bytes:
        """Download document with retry mechanism"""
        for attempt in range(max_retries):
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.get(url, timeout=30) as response:
                        if response.status == 200:
                            content = await response.read()
                            if len(content) > 100 * 1024 * 1024:  # 100MB limit
                                raise ValueError(f"Document too large: {len(content) / 1024 / 1024:.2f}MB")
                            return content
                        else:
                            raise aiohttp.ClientError(f"HTTP {response.status}")
            except Exception as e:
                if attempt == max_retries - 1:
                    raise e
                await asyncio.sleep(2 ** attempt)  # Exponential backoff
    
    async def process_pdf_advanced(self, content: bytes) -> Dict[str, Any]:
        """Advanced PDF processing with image and table extraction"""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            logger.info(f"PDF opened successfully. Pages: {len(doc)}")
        except Exception as e:
            logger.error(f"Failed to open PDF: {e}")
            # Fallback: try to extract text from raw content
            return {
                'text': content.decode('utf-8', errors='ignore'),
                'images': [],
                'tables': [],
                'sections': [],
                'metadata': {}
            }
            
        extracted_data = {
            'text': '',
            'images': [],
            'tables': [],
            'sections': [],
            'metadata': {}
        }
        
        try:
            # Simple, reliable text extraction
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    extracted_data['text'] += page_text + '\n\n'
                    logger.info(f"Page {page_num + 1}: Extracted {len(page_text)} characters")
                
                # Extract images with OCR
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        if pix.n - pix.alpha < 4:  # Valid image
                            img_data = pix.tobytes("png")
                            ocr_text = await self.extract_text_from_image(img_data)
                            
                            extracted_data['images'].append({
                                'page': page_num + 1,
                                'ocr_text': ocr_text,
                                'confidence': self.estimate_ocr_confidence(ocr_text),
                                'size': (pix.width, pix.height)
                            })
                        pix = None
                    except Exception as e:
                        logger.warning(f"Image extraction failed on page {page_num}: {e}")
                        continue
                
                # Extract tables
                tables = page.find_tables()
                for table in tables:
                    try:
                        table_data = table.extract()
                        extracted_data['tables'].append({
                            'page': page_num + 1,
                            'data': table_data,
                            'bbox': table.bbox
                        })
                    except Exception as e:
                        logger.warning(f"Table extraction failed on page {page_num}: {e}")
                        continue
            
            # Analyze document structure
            extracted_data['sections'] = self.analyze_document_structure(extracted_data['text'])
            
        finally:
            doc.close()
        
        return extracted_data
    
    async def process_docx_advanced(self, content: bytes) -> Dict[str, Any]:
        """Advanced DOCX processing with style and table extraction"""
        doc = docx.Document(io.BytesIO(content))
        extracted_data = {
            'text': '',
            'images': [],
            'tables': [],
            'sections': [],
            'metadata': {}
        }
        
        # Extract text with formatting
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_data['text'] += paragraph.text + '\n\n'
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            extracted_data['tables'].append({
                'data': table_data,
                'page': 1  # DOCX doesn't have page info
            })
        
        # Analyze structure
        extracted_data['sections'] = self.analyze_document_structure(extracted_data['text'])
        
        return extracted_data
    
    async def process_email_advanced(self, content: bytes) -> Dict[str, Any]:
        """Advanced email processing with attachment handling"""
        msg = email.message_from_bytes(content, policy=policy.default)
        extracted_data = {
            'text': '',
            'images': [],
            'tables': [],
            'sections': [],
            'metadata': {
                'subject': msg.get('subject', ''),
                'from': msg.get('from', ''),
                'to': msg.get('to', ''),
                'date': msg.get('date', '')
            }
        }
        
        # Extract email body
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    extracted_data['text'] += part.get_content() + '\n\n'
        else:
            extracted_data['text'] = msg.get_content()
        
        # Analyze structure
        extracted_data['sections'] = self.analyze_document_structure(extracted_data['text'])
        
        return extracted_data
    
    async def process_generic_text(self, content: bytes) -> Dict[str, Any]:
        """Process generic text content"""
        text = content.decode('utf-8', errors='ignore')
        return {
            'text': text,
            'images': [],
            'tables': [],
            'sections': self.analyze_document_structure(text),
            'metadata': {}
        }
    
    def extract_structured_text(self, text_dict: Dict) -> str:
        """Extract structured text from PDF text dictionary"""
        text = ""
        for block in text_dict.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    for span in line.get("spans", []):
                        text += span.get("text", "")
                    text += "\n"
        return text
    
    async def extract_text_from_image(self, img_data: bytes) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(io.BytesIO(img_data))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return ""
    
    def estimate_ocr_confidence(self, text: str) -> float:
        """Estimate OCR confidence based on text characteristics"""
        if not text:
            return 0.0
        
        # Simple heuristic: longer text with proper spacing is more confident
        confidence = min(1.0, len(text) / 100.0)
        
        # Penalize for unusual characters
        unusual_chars = len(re.findall(r'[^a-zA-Z0-9\s.,!?;:]', text))
        if len(text) > 0:
            confidence *= (1 - unusual_chars / len(text))
        
        return max(0.0, confidence)
    
    def analyze_document_structure(self, text: str) -> List[Dict[str, Any]]:
        """Analyze document structure and extract sections"""
        sections = []
        lines = text.split('\n')
        current_section = {"title": "", "content": "", "level": 0}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect headers (simple heuristic)
            if len(line) < 100 and (line.isupper() or line.endswith(':') or re.match(r'^[0-9]+\.', line)):
                if current_section["content"]:
                    sections.append(current_section.copy())
                
                current_section = {
                    "title": line,
                    "content": "",
                    "level": self.determine_header_level(line)
                }
            else:
                current_section["content"] += line + "\n"
        
        if current_section["content"]:
            sections.append(current_section)
        
        return sections
    
    def determine_header_level(self, line: str) -> int:
        """Determine header level based on formatting"""
        if line.isupper():
            return 1
        elif line.endswith(':'):
            return 2
        elif re.match(r'^[0-9]+\.', line):
            return 3
        else:
            return 0
    
    def validate_extraction_quality(self, result: Dict[str, Any]) -> float:
        """Validate extraction quality and return score"""
        score = 0.0
        total_checks = 0
        
        # Check text length
        if result.get('text'):
            text_length = len(result['text'])
            if text_length >= self.quality_thresholds['min_text_length']:
                score += 1.0
            total_checks += 1
        
        # Check sections
        if result.get('sections'):
            if len(result['sections']) > 0:
                score += 1.0
            total_checks += 1
        
        # Check images
        if result.get('images'):
            valid_images = len([img for img in result['images'] if img.get('ocr_text')])
            if valid_images > 0:
                score += 1.0
            total_checks += 1
        
        # Check tables
        if result.get('tables'):
            if len(result['tables']) > 0:
                score += 1.0
            total_checks += 1
        
        return score / total_checks if total_checks > 0 else 0.0
    
    async def enhance_extraction(self, result: Dict[str, Any], content: bytes) -> Dict[str, Any]:
        """Enhance extraction when quality is low"""
        logger.info("Enhancing document extraction")
        
        # Try alternative extraction methods
        if not result.get('text') or len(result['text']) < 100:
            # Try OCR on the entire document
            try:
                images = convert_from_bytes(content)
                ocr_text = ""
                for image in images:
                    ocr_text += pytesseract.image_to_string(image) + "\n"
                result['text'] = ocr_text
            except Exception as e:
                logger.warning(f"OCR enhancement failed: {e}")
        
        return result
    
    def calculate_extraction_confidence(self, result: Dict[str, Any]) -> float:
        """Calculate overall extraction confidence"""
        confidence = 0.0
        factors = 0
        
        # Text confidence
        if result.get('text'):
            text_confidence = min(1.0, len(result['text']) / 1000.0)
            confidence += text_confidence
            factors += 1
        
        # Image confidence
        if result.get('images'):
            avg_img_confidence = sum(img.get('confidence', 0) for img in result['images']) / len(result['images'])
            confidence += avg_img_confidence
            factors += 1
        
        # Structure confidence
        if result.get('sections'):
            structure_confidence = min(1.0, len(result['sections']) / 10.0)
            confidence += structure_confidence
            factors += 1
        
        return confidence / factors if factors > 0 else 0.5
    
    def detect_document_format(self, url: str, content: bytes) -> str:
        """Detect document format"""
        if url.lower().endswith('.pdf'):
            return 'pdf'
        elif url.lower().endswith('.docx'):
            return 'docx'
        elif url.lower().endswith('.eml'):
            return 'email'
        else:
            return 'text'
    
    def handle_processing_error(self, error: Exception, url: str) -> Dict[str, Any]:
        """Handle processing errors gracefully"""
        logger.error(f"Document processing failed for {url}: {error}")
        
        return {
            'text': f"Error processing document: {str(error)}",
            'images': [],
            'tables': [],
            'sections': [],
            'metadata': {
                'error': True,
                'error_type': type(error).__name__,
                'error_message': str(error),
                'url': url
            },
            'processing_metadata': {
                'quality_score': 0.0,
                'extraction_confidence': 0.0,
                'error': True
            }
        } 

    def process_text(self, text: str) -> list:
        """Simple chunking for demonstration: splits text by double newlines."""
        chunks = []
        for i, paragraph in enumerate(text.split('\n\n')):
            clean = paragraph.strip()
            if clean:
                chunks.append({'id': f'chunk_{i}', 'text': clean})
        return chunks 